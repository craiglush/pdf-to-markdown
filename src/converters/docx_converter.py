"""
DOCX to Markdown converter using pypandoc and mammoth.

.. deprecated:: 2.0.0
    This converter is deprecated in favor of MarkItDownConverter.
    Use MarkItDownConverter for better DOCX support and unified multi-format conversion.
"""

import warnings
import io
import logging
import re
import tempfile
from pathlib import Path
from typing import List, Optional

from PIL import Image as PILImage

from pdf2markdown.converters.document_converter import DocumentConverter

# Deprecation warning
warnings.warn(
    "DOCXConverter is deprecated as of version 2.0.0. Use MarkItDownConverter instead.",
    DeprecationWarning,
    stacklevel=2
)
from pdf2markdown.core.config import Config
from pdf2markdown.core.models import (
    ConversionMetadata,
    ConversionResult,
    ExtractedImage,
    ExtractedTable,
)

logger = logging.getLogger(__name__)


class DOCXConverter(DocumentConverter):
    """Convert DOCX documents to Markdown.

    Uses pypandoc (requires pandoc) as primary converter with mammoth as fallback.
    Extracts images and metadata using python-docx.
    """

    def __init__(self, config: Optional[Config] = None):
        """Initialize the DOCX converter.

        Args:
            config: Conversion configuration
        """
        self.config = config or Config()
        self._pypandoc_available = self._check_pypandoc()
        self._mammoth_available = self._check_mammoth()

    def convert(self, file_path: Path) -> ConversionResult:
        """Convert a DOCX file to Markdown.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ConversionResult with markdown, images, and metadata

        Raises:
            FileNotFoundError: If the DOCX file doesn't exist
            ValueError: If the file is not valid DOCX or conversion fails
        """
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        logger.info(f"Converting DOCX file: {file_path}")

        # Extract metadata first
        metadata = self._extract_metadata(file_path)

        # Extract images
        images = self._extract_images(file_path)

        # Convert to Markdown (try pypandoc first, then mammoth)
        markdown = None
        conversion_method = None

        if self._pypandoc_available:
            try:
                markdown = self._convert_with_pypandoc(file_path)
                conversion_method = "pypandoc"
                logger.info("Conversion successful with pypandoc")
            except Exception as e:
                logger.warning(f"pypandoc conversion failed: {e}, trying mammoth...")

        if markdown is None and self._mammoth_available:
            try:
                markdown = self._convert_with_mammoth(file_path)
                conversion_method = "mammoth"
                logger.info("Conversion successful with mammoth")
            except Exception as e:
                logger.error(f"mammoth conversion failed: {e}")

        if markdown is None:
            raise ValueError(
                "DOCX conversion failed. Install pypandoc/pandoc or mammoth: "
                "pip install pypandoc mammoth"
            )

        # Post-process markdown
        markdown = self._postprocess_markdown(markdown)

        # Extract tables (basic implementation)
        tables = self._extract_tables(markdown)

        # Update metadata with converter info
        metadata.converter_name = f"DOCX Converter ({conversion_method})"

        logger.info(f"Conversion complete: {len(markdown)} chars, "
                   f"{len(images)} images, {len(tables)} tables")

        return ConversionResult(
            markdown=markdown,
            images=images,
            tables=tables,
            metadata=metadata,
        )

    def supports_ocr(self) -> bool:
        """DOCX converter doesn't need OCR."""
        return False

    def get_name(self) -> str:
        """Get converter name."""
        return "DOCX Converter"

    def is_available(self) -> bool:
        """Check if required dependencies are available."""
        # At least one converter must be available
        return self._pypandoc_available or self._mammoth_available

    def get_supported_extensions(self) -> List[str]:
        """Get supported file extensions."""
        return ['.docx']

    def _check_pypandoc(self) -> bool:
        """Check if pypandoc and pandoc are available."""
        try:
            import pypandoc
            # Check if pandoc is installed
            pypandoc.get_pandoc_version()
            return True
        except (ImportError, OSError):
            return False

    def _check_mammoth(self) -> bool:
        """Check if mammoth is available."""
        try:
            import mammoth  # noqa: F401
            return True
        except ImportError:
            return False

    def _extract_metadata(self, file_path: Path) -> ConversionMetadata:
        """Extract metadata from DOCX document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Conversion metadata
        """
        try:
            import docx

            doc = docx.Document(file_path)
            core_props = doc.core_properties

            # Extract metadata
            title = core_props.title
            author = core_props.author
            subject = core_props.subject

            # Count paragraphs as "pages" (rough estimate)
            paragraph_count = len(doc.paragraphs)
            # Estimate pages (assuming ~25 paragraphs per page)
            page_count = max(1, paragraph_count // 25)

            return ConversionMetadata(
                title=title,
                author=author,
                subject=subject,
                page_count=page_count,
                converter_name=self.get_name(),
            )

        except ImportError:
            logger.warning("python-docx not installed, metadata extraction limited")
            return ConversionMetadata(
                title=None,
                author=None,
                subject=None,
                page_count=1,
                converter_name=self.get_name(),
            )
        except Exception as e:
            logger.error(f"Error extracting metadata: {e}")
            return ConversionMetadata(
                title=None,
                author=None,
                subject=None,
                page_count=1,
                converter_name=self.get_name(),
            )

    def _extract_images(self, file_path: Path) -> List[ExtractedImage]:
        """Extract images from DOCX document.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of extracted images
        """
        images = []

        try:
            import docx
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            doc = docx.Document(file_path)

            # Iterate through document parts to find images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_part = rel.target_part
                        image_data = image_part.blob

                        # Get image metadata
                        img = PILImage.open(io.BytesIO(image_data))
                        width, height = img.size
                        img_format = img.format or "unknown"

                        # Generate image name
                        img_name = Path(rel.target_ref).name

                        images.append(ExtractedImage(
                            name=img_name,
                            data=image_data,
                            width=width,
                            height=height,
                            page=0,  # DOCX doesn't have pages
                            format=img_format.lower(),
                            alt_text="",
                        ))

                    except Exception as e:
                        logger.warning(f"Error extracting image {rel.target_ref}: {e}")

        except ImportError:
            logger.warning("python-docx not installed, image extraction skipped")
        except Exception as e:
            logger.error(f"Error extracting images: {e}")

        return images

    def _convert_with_pypandoc(self, file_path: Path) -> str:
        """Convert DOCX to Markdown using pypandoc.

        Args:
            file_path: Path to DOCX file

        Returns:
            Markdown string

        Raises:
            Exception: If conversion fails
        """
        import pypandoc

        # Pypandoc options
        extra_args = [
            '--wrap=none',  # Don't wrap lines
            '--extract-media=.',  # Extract images to current dir
        ]

        # Add options based on config
        if self.config.docx_include_comments:
            extra_args.append('--track-changes=all')

        # Convert
        markdown = pypandoc.convert_file(
            str(file_path),
            'markdown',
            format='docx',
            extra_args=extra_args,
        )

        return markdown

    def _convert_with_mammoth(self, file_path: Path) -> str:
        """Convert DOCX to Markdown using mammoth.

        Args:
            file_path: Path to DOCX file

        Returns:
            Markdown string

        Raises:
            Exception: If conversion fails
        """
        import mammoth

        # Custom style map for better markdown conversion
        style_map = """
        p[style-name='Heading 1'] => h1:fresh
        p[style-name='Heading 2'] => h2:fresh
        p[style-name='Heading 3'] => h3:fresh
        p[style-name='Heading 4'] => h4:fresh
        p[style-name='Heading 5'] => h5:fresh
        p[style-name='Heading 6'] => h6:fresh
        """

        with open(file_path, 'rb') as docx_file:
            result = mammoth.convert_to_markdown(
                docx_file,
                style_map=style_map,
            )

        if result.messages:
            for message in result.messages:
                logger.debug(f"Mammoth message: {message}")

        return result.value

    def _extract_tables(self, markdown: str) -> List[ExtractedTable]:
        """Extract tables from markdown.

        Args:
            markdown: Markdown string

        Returns:
            List of extracted tables
        """
        tables = []

        # Find all markdown tables
        table_pattern = r'(\|[^\n]+\|\n)+\|[-:\s|]+\|(\n\|[^\n]+\|)+'
        matches = re.finditer(table_pattern, markdown)

        for idx, match in enumerate(matches):
            table_markdown = match.group(0)

            # Count rows and columns
            lines = table_markdown.strip().split('\n')
            row_count = len([l for l in lines if not re.match(r'\|[-:\s|]+\|', l)])
            col_count = len(lines[0].split('|')) - 2 if lines else 0

            tables.append(ExtractedTable(
                markdown=table_markdown,
                page=0,
                caption=None,
                row_count=row_count,
                col_count=col_count,
            ))

        return tables

    def _postprocess_markdown(self, markdown: str) -> str:
        """Post-process generated Markdown.

        Args:
            markdown: Raw markdown string

        Returns:
            Cleaned markdown string
        """
        # Remove excessive blank lines (more than 2 consecutive)
        markdown = re.sub(r'\n{3,}', '\n\n', markdown)

        # Clean up spacing around headings
        markdown = re.sub(r'\n(#{1,6} .+)\n', r'\n\n\1\n\n', markdown)

        # Ensure proper spacing around code blocks
        markdown = re.sub(r'```(\w+)?\n', r'\n```\1\n', markdown)
        markdown = re.sub(r'\n```\n', r'\n```\n\n', markdown)

        # Clean up list formatting
        markdown = re.sub(r'\n([\*\-\+\d]+\. .+)\n', r'\n\1\n', markdown)

        # Remove leading/trailing whitespace
        markdown = markdown.strip()

        return markdown
